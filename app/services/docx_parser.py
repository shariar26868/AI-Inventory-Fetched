import logging
from typing import List, Dict, Any
from docx import Document

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse a .docx file. If tables are present, convert table rows to dicts
    using the first row as headers. Otherwise, return non-empty paragraphs
    as raw_text entries.
    """
    try:
        doc = Document(file_path)
        results: List[Dict[str, Any]] = []

        # Parse tables first
        for t_idx, table in enumerate(doc.tables, start=1):
            if len(table.rows) < 1:
                continue
            # header texts
            headers = [cell.text.strip() or f"col_{i}" for i, cell in enumerate(table.rows[0].cells)]
            for row in table.rows[1:]:
                row_dict: Dict[str, Any] = {}
                for i, cell in enumerate(row.cells):
                    key = headers[i] if i < len(headers) else f"col_{i}"
                    row_dict[key] = cell.text.strip() if cell.text else None
                row_dict["_table"] = t_idx
                row_dict["_source"] = "docx"
                results.append(row_dict)

        # Fallback: paragraphs
        if not results:
            for p_idx, para in enumerate(doc.paragraphs, start=1):
                text = para.text.strip()
                if not text:
                    continue
                results.append({
                    "raw_text": text,
                    "_para": p_idx,
                    "_source": "docx",
                })

        logger.info(f"DOCX parsed: {len(results)} entries from {file_path}")
        return results

    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        raise
